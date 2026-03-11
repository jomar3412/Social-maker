"""
Document Generator

Generates styled DOCX documents for VEO prompts and scene breakdowns.
Creates professional, readable documents while keeping .md for machine parsing.
"""

from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _create_styled_document() -> "Document":
    """Create a new Document with custom styles."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    return doc


def _add_title(doc: "Document", title: str, subtitle: str = None):
    """Add a styled title to the document."""
    # Main title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle if provided
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.runs[0]
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Spacing


def _add_info_box(doc: "Document", label: str, content: str, is_code: bool = False):
    """Add a labeled info box."""
    para = doc.add_paragraph()

    # Label (bold)
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True

    # Content
    if is_code:
        # Add as monospace
        content_run = para.add_run(content)
        content_run.font.name = 'Consolas'
        content_run.font.size = Pt(10)
    else:
        para.add_run(content)


def _add_code_block(doc: "Document", code: str):
    """Add a styled code block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)

    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(50, 50, 50)


def _add_scene_section(
    doc: "Document",
    scene_num: int,
    voiceover: str,
    prompt: str,
    clip_file: str = None,
):
    """Add a scene section with voiceover and VEO prompt."""
    # Scene header
    header_text = f"Scene {scene_num}"
    if clip_file:
        header_text += f"  ({clip_file})"
    doc.add_heading(header_text, level=2)

    # Voiceover
    vo_para = doc.add_paragraph()
    vo_label = vo_para.add_run("Voiceover: ")
    vo_label.bold = True
    vo_para.add_run(f'"{voiceover}"')

    doc.add_paragraph()  # Spacing

    # VEO Prompt label
    prompt_label = doc.add_paragraph()
    label_run = prompt_label.add_run("VEO Prompt:")
    label_run.bold = True

    # VEO Prompt content (styled as code block)
    _add_code_block(doc, prompt)

    # Divider
    doc.add_paragraph("_" * 60)


def generate_veo_prompts_docx(
    output_path: Path,
    video_id: str,
    version: int,
    prompts: List[Dict],
    style_lock: str,
    primary_subject: str,
) -> Optional[Path]:
    """
    Generate a styled DOCX document for VEO prompts.

    Args:
        output_path: Path for output file (will use .docx extension)
        video_id: The video ID
        version: Version number
        prompts: List of {"scene": 1, "voiceover": "...", "prompt": "..."}
        style_lock: The style lock used for all scenes
        primary_subject: The primary subject identified

    Returns:
        Path to the saved DOCX file, or None if docx not available
    """
    if not DOCX_AVAILABLE:
        print("  Warning: python-docx not installed. Run: pip install python-docx")
        return None

    doc = _create_styled_document()

    # Title
    _add_title(
        doc,
        f"VEO Prompts",
        f"{video_id} - Version {version}"
    )

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Style Lock Section
    doc.add_heading("Style Lock (Applied to ALL Scenes)", level=1)
    _add_code_block(doc, style_lock)

    # Primary Subject
    _add_info_box(doc, "Primary Subject", primary_subject)

    doc.add_paragraph()
    doc.add_paragraph("_" * 60)

    # Scene-by-Scene Prompts
    doc.add_heading("Scene-by-Scene Prompts", level=1)

    for p in prompts:
        scene_num = p.get("scene", 0)
        voiceover = p.get("voiceover", "")
        prompt = p.get("prompt", "")
        clip_file = f"scene_{scene_num:02d}.mp4"

        _add_scene_section(doc, scene_num, voiceover, prompt, clip_file)

    # Save
    docx_path = output_path.with_suffix('.docx')
    doc.save(str(docx_path))

    return docx_path


def generate_scene_breakdown_docx(
    output_path: Path,
    scenes: List,
    video_id: str = None,
    total_duration: float = None,
) -> Optional[Path]:
    """
    Generate a styled DOCX document for scene breakdown.

    Args:
        output_path: Path for output file (will use .docx extension)
        scenes: List of Scene objects or dicts
        video_id: Optional video ID for title
        total_duration: Optional total duration

    Returns:
        Path to the saved DOCX file, or None if docx not available
    """
    if not DOCX_AVAILABLE:
        print("  Warning: python-docx not installed. Run: pip install python-docx")
        return None

    doc = _create_styled_document()

    # Title
    title = "Scene Breakdown"
    subtitle = video_id if video_id else None
    _add_title(doc, title, subtitle)

    # Summary info
    summary = doc.add_paragraph()
    summary.add_run(f"Total Scenes: {len(scenes)}")
    if total_duration:
        summary.add_run(f"  |  Duration: {total_duration:.1f}s")
    summary.add_run(f"  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_paragraph("_" * 60)

    # Each scene
    for scene in scenes:
        # Handle both Scene objects and dicts
        if hasattr(scene, 'scene_number'):
            scene_num = scene.scene_number
            duration = scene.duration
            start = scene.start_time
            end = scene.end_time
            on_screen = scene.on_screen_text
            voiceover = scene.voiceover_text
            visual = scene.visual_direction
            animation = scene.text_animation.value if hasattr(scene.text_animation, 'value') else str(scene.text_animation)
            keywords = scene.keywords
        else:
            scene_num = scene.get('scene_number', scene.get('scene', 0))
            duration = scene.get('duration', scene.get('end', 0) - scene.get('start', 0))
            start = scene.get('start_time', scene.get('start', 0))
            end = scene.get('end_time', scene.get('end', 0))
            on_screen = scene.get('on_screen_text', '')
            voiceover = scene.get('voiceover_text', '')
            visual = scene.get('visual_direction', '')
            animation = scene.get('text_animation', '')
            keywords = scene.get('keywords', [])

        # Scene header
        doc.add_heading(f"Scene {scene_num}", level=2)

        # Timing
        timing = doc.add_paragraph()
        timing_run = timing.add_run("Duration: ")
        timing_run.bold = True
        timing.add_run(f"{duration:.1f}s ({start:.1f}s - {end:.1f}s)")

        # On-screen text
        if on_screen:
            text_para = doc.add_paragraph()
            text_label = text_para.add_run("On-Screen Text: ")
            text_label.bold = True
            text_para.add_run(f'"{on_screen}"')

        # Voiceover (truncated if long)
        if voiceover:
            vo_para = doc.add_paragraph()
            vo_label = vo_para.add_run("Voiceover: ")
            vo_label.bold = True
            display_vo = voiceover[:100] + "..." if len(voiceover) > 100 else voiceover
            vo_para.add_run(display_vo)

        # Visual Direction (this is the VEO prompt)
        if visual:
            doc.add_paragraph()
            visual_label = doc.add_paragraph()
            label_run = visual_label.add_run("Visual Direction / VEO Prompt:")
            label_run.bold = True
            _add_code_block(doc, visual)

        # Animation
        if animation:
            anim_para = doc.add_paragraph()
            anim_label = anim_para.add_run("Text Animation: ")
            anim_label.bold = True
            anim_para.add_run(animation)

        # Keywords
        if keywords:
            kw_para = doc.add_paragraph()
            kw_label = kw_para.add_run("Keywords: ")
            kw_label.bold = True
            kw_para.add_run(", ".join(keywords))

        doc.add_paragraph("_" * 60)

    # Save
    docx_path = output_path.with_suffix('.docx')
    doc.save(str(docx_path))

    return docx_path


if __name__ == "__main__":
    # Test document generation
    print("Testing document generation...")

    if not DOCX_AVAILABLE:
        print("python-docx not installed. Run: pip install python-docx")
    else:
        # Test VEO prompts document
        test_prompts = [
            {
                "scene": 1,
                "voiceover": "Butterflies taste with their feet.",
                "prompt": "3D animated monarch butterfly, macro documentary style, tasting with feet on flower",
            },
            {
                "scene": 2,
                "voiceover": "Their feet have special sensors.",
                "prompt": "3D animated monarch butterfly, macro documentary style, showing foot sensors in detail",
            },
        ]

        output = Path("test_veo_prompts.docx")
        result = generate_veo_prompts_docx(
            output,
            video_id="VIDEO-BUTTERFLY-20260212-001",
            version=1,
            prompts=test_prompts,
            style_lock="3D animated monarch butterfly, macro documentary style, soft natural lighting",
            primary_subject="monarch butterfly",
        )

        if result:
            print(f"Created: {result}")

        print("Done!")
